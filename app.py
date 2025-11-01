# app.py — RAG para Revisão Bibliográfica (Streamlit + FAISS + Sentence-Transformers + OpenAI/Gemini)
# Execução: streamlit run app.py

from __future__ import annotations
import os
import io
import re
import collections
from dataclasses import dataclass
from typing import List, Dict, Tuple, Optional, Set

import streamlit as st
import fitz  # PyMuPDF
from sentence_transformers import SentenceTransformer
import faiss

# ====== OpenAI (principal) ======
USE_OPENAI = False
try:
    from openai import OpenAI  # pip install openai>=1.40.0
    USE_OPENAI = True
except Exception:
    USE_OPENAI = False

# ====== Gemini (fallback) ======
USE_GEMINI = False
try:
    import google.generativeai as genai  # pip install google-generativeai>=0.7.2
    USE_GEMINI = True
except Exception:
    USE_GEMINI = False

# ====== Exportação para Word (opcional) ======
try:
    from docx import Document  # pip install python-docx
    HAS_DOCX = True
except Exception:
    HAS_DOCX = False

DEFAULT_EMB = os.getenv("EMBEDDING_MODEL", "sentence-transformers/all-MiniLM-L6-v2")

# =========================
# Estruturas de dados
# =========================

@dataclass
class Chunk:
    doc_id: str
    chunk_id: str
    text: str
    page_start: int
    page_end: int
    title: str

@dataclass
class RefEntry:
    raw: str
    authors: List[str]           # ["Sobrenome, A. A.", "Sobrenome, B. B.", ...]
    year: str                    # "2021" ou "2021a"
    title: str                   # artigo/capítulo
    container: Optional[str]     # periódico / livro / conferência
    volume: Optional[str]
    issue: Optional[str]
    pages: Optional[str]
    doi: Optional[str]
    url: Optional[str]

    def first_author_key(self) -> Optional[str]:
        return (self.authors[0].split(",")[0].lower()) if self.authors else None

    def citation_key(self) -> Optional[Tuple[str, str]]:
        fa = self.first_author_key()
        return (fa, self.year.lower()) if (fa and self.year) else None

# =========================
# Utilidades gerais
# =========================

def clean_text(s: str) -> str:
    s = re.sub(r"\s+", " ", s)
    return s.strip()

def canonicalize(text: str) -> str:
    return re.sub(r"\s+", " ", text.strip().lower())

def read_pdf(path_or_bytes: bytes | str) -> Tuple[List[str], str]:
    """Lê PDF e retorna lista de textos por página e um 'título' (nome do arquivo)."""
    if isinstance(path_or_bytes, (bytes, bytearray)):
        doc = fitz.open(stream=path_or_bytes, filetype="pdf")
        title = "uploaded.pdf"
    else:
        doc = fitz.open(path_or_bytes)
        title = os.path.basename(path_or_bytes)
    pages = [p.get_text("text") for p in doc]
    return pages, title

# =========================
# Chunking e índice vetorial
# =========================

def chunk_pages(
    pages: List[str],
    title: str,
    doc_id: str,
    chunk_size: int = 1200,
    overlap: int = 200
) -> List[Chunk]:
    """Cria chunks deslizantes preservando mapeamento aproximado para páginas."""
    joined = "\n".join(pages)
    tokens = joined.split()

    # mapa token_index -> índice de página
    page_offsets = []
    tok_count = 0
    for i, page in enumerate(pages):
        n = len(page.split())
        page_offsets.append((tok_count, i))
        tok_count += n

    def token_index_to_page(idx: int) -> int:
        best = 0
        for (start_tok, page_idx) in page_offsets:
            if idx >= start_tok:
                best = page_idx
            else:
                break
        return best

    chunks: List[Chunk] = []
    start = 0
    cid = 0
    while start < len(tokens):
        end = min(len(tokens), start + chunk_size)
        text = clean_text(" ".join(tokens[start:end]))
        p_start = token_index_to_page(start)
        p_end = token_index_to_page(end)
        chunks.append(
            Chunk(
                doc_id=doc_id,
                chunk_id=f"{doc_id}-{cid}",
                text=text,
                page_start=p_start,
                page_end=p_end,
                title=title,
            )
        )
        cid += 1
        if end == len(tokens):
            break
        start = max(0, end - overlap)
    return chunks

class VectorIndex:
    def __init__(self, model_name: str = DEFAULT_EMB):
        self.model_name = model_name
        self.emb_model = SentenceTransformer(model_name)
        self.index: Optional[faiss.Index] = None
        self.vectors = None
        self.meta: List[Chunk] = []

    def add(self, chunks: List[Chunk]):
        embs = self.emb_model.encode(
            [c.text for c in chunks],
            convert_to_numpy=True,
            normalize_embeddings=True
        )
        if self.index is None:
            self.index = faiss.IndexFlatIP(embs.shape[1])
            self.vectors = embs
        else:
            import numpy as np
            self.vectors = np.vstack([self.vectors, embs])
        self.meta.extend(chunks)
        self.index.add(embs)

    def search(self, query: str, k: int = 20) -> List[Tuple[Chunk, float]]:
        q = self.emb_model.encode([query], convert_to_numpy=True, normalize_embeddings=True)
        D, I = self.index.search(q, k)
        results: List[Tuple[Chunk, float]] = []
        for idx, score in zip(I[0], D[0]):
            if idx == -1:
                continue
            results.append((self.meta[idx], float(score)))
        return results

def diversify_by_doc(
    hits: List[Tuple[Chunk, float]],
    per_doc: int = 3,
    max_total: int = 12
) -> List[Chunk]:
    """Round-robin por documento: limita nº de chunks por paper e total."""
    bydoc: Dict[str, List[Tuple[Chunk, float]]] = {}
    for ch, sc in hits:
        bydoc.setdefault(ch.doc_id, []).append((ch, sc))
    for d in bydoc:
        bydoc[d].sort(key=lambda x: x[1], reverse=True)

    out: List[Chunk] = []
    round_i = 0
    while len(out) < max_total:
        added = 0
        for d in sorted(bydoc.keys()):
            if round_i < len(bydoc[d]) and round_i < per_doc:
                out.append(bydoc[d][round_i][0])
                added += 1
                if len(out) >= max_total:
                    break
        if added == 0:
            break
        round_i += 1
    return out

def build_citation(ch: Chunk) -> str:
    base = os.path.splitext(os.path.basename(ch.title))[0]
    return f"[{base}:{ch.page_start+1}]"

# =========================
# Extração de referências dos PDFs
# =========================

REF_HEADERS = [
    "references", "referências", "bibliography", "bibliografia", "referencias"
]
CITATION_PATTERNS = [
    r"\(([A-Z][A-Za-z\-']+),\s*(\d{4}[a-z]?)\)",
    r"\(([A-Z][A-Za-z\-']+)\s+et al\.,\s*(\d{4}[a-z]?)\)",
    r"\(([A-Z][A-Za-z\-']+)\s*&\s*([A-Z][A-Za-z\-']+),\s*(\d{4}[a-z]?)\)",
    r"\(([A-Z][A-Za-z\-']+),\s*([A-Z][A-Za-z\-']+),\s*(\d{4}[a-z]?)\)",
]
DOI_RX = re.compile(r"\b10\.\d{4,9}/\S+\b", flags=re.I)
URL_RX = re.compile(r"https?://\S+", flags=re.I)
AUTHORS_YEAR_RX = re.compile(r"^([A-Z][^()]+?)\s*\((\d{4}[a-z]?)\)\.?\s*", flags=re.U)
PAGES_RX = re.compile(r"\b(pp?\.)\s*([\d\-–—]+)\b", flags=re.I)
VOL_ISSUE_RX = re.compile(r"(\d+)\s*\((\d+)\)")
VOL_ONLY_RX = re.compile(r"\b(\d+)\b")

def _looks_like_ref_header(line: str) -> bool:
    s = line.strip().lower()
    return any(s.startswith(h) for h in REF_HEADERS)

def _split_nonempty_lines(text: str) -> List[str]:
    return [ln.strip() for ln in text.splitlines() if ln.strip()]

def extract_references_section_lines(pages: List[str]) -> List[str]:
    if not pages:
        return []
    start_page = max(0, int(0.6 * (len(pages)-1)))  # parte final
    tail = "\n".join(pages[start_page:])
    lines = _split_nonempty_lines(tail)
    refs_started = False
    ref_lines: List[str] = []
    for ln in lines:
        if not refs_started and _looks_like_ref_header(ln):
            refs_started = True
            continue
        if refs_started:
            if ln.lower().startswith(("appendix", "apêndice", "apendice")):
                break
            ref_lines.append(ln)
    return ref_lines

def _split_reference_blocks(ref_lines: List[str]) -> List[str]:
    blocks, buf = [], []
    for ln in ref_lines:
        buf.append(ln)
        if ln.endswith(".") or ln.endswith(").") or ln.endswith("]."):
            blocks.append(" ".join(buf))
            buf = []
    if buf:
        blocks.append(" ".join(buf))
    return blocks

def _parse_authors_list(authors_raw: str) -> List[str]:
    tmp = re.split(r"\s*&\s*|\s+and\s+|;\s*", authors_raw)
    out: List[str] = []
    for part in tmp:
        part = part.strip().rstrip(",")
        if "," in part:
            out.append(part)
    return out

def parse_reference_block(block: str) -> RefEntry:
    raw = block.strip()

    doi = None
    mdoi = DOI_RX.search(raw)
    if mdoi:
        doi = mdoi.group(0).rstrip(").,;")
    url = None
    murl = URL_RX.search(raw)
    if murl:
        url = murl.group(0).rstrip(").,;")

    authors, year = [], ""
    m = AUTHORS_YEAR_RX.search(raw)
    rest = raw
    if m:
        authors_raw, year = m.group(1), m.group(2)
        authors = _parse_authors_list(authors_raw)
        rest = raw[m.end():].strip()

    title, container, volume, issue, pages = "", None, None, None, None
    parts = [p.strip() for p in rest.split(".") if p.strip()]
    if parts:
        title = parts[0]
        tail = ". ".join(parts[1:])
        container = None
        if "," in tail:
            container = tail.split(",")[0].strip()
        mi = VOL_ISSUE_RX.search(tail)
        if mi:
            volume, issue = mi.group(1), mi.group(2)
        else:
            mv = VOL_ONLY_RX.search(tail)
            if mv:
                volume = mv.group(1)
        mp = PAGES_RX.search(tail)
        if mp:
            pages = mp.group(2)

    return RefEntry(
        raw=raw, authors=authors, year=year, title=title,
        container=container, volume=volume, issue=issue,
        pages=pages, doi=doi, url=url
    )

def format_apa(entry: RefEntry) -> str:
    # Autores (APA básico)
    if entry.authors:
        if len(entry.authors) == 1:
            a = entry.authors[0]
        elif len(entry.authors) == 2:
            a = f"{entry.authors[0]}, & {entry.authors[1]}"
        else:
            a = ", ".join(entry.authors[:-1]) + f", & {entry.authors[-1]}"
    else:
        a = "[Autor não identificado]"

    y = f"({entry.year})." if entry.year else "(s.d.)."
    t = (entry.title.rstrip(".") + ".") if entry.title else "[Sem título]."

    cont = (entry.container.rstrip(".") + ",") if entry.container else ""
    vol_issue = ""
    if entry.volume and entry.issue:
        vol_issue = f" {entry.volume}({entry.issue}),"
    elif entry.volume:
        vol_issue = f" {entry.volume},"

    pg = f" {entry.pages}." if entry.pages else ""
    tail = ""
    if entry.doi:
        tail = f" https://doi.org/{entry.doi.split('10.', 1)[-1]}" if entry.doi.lower().startswith("10.") else f" {entry.doi}"
    elif entry.url:
        tail = f" {entry.url}"

    # Montagem
    core = f"{a} {y} {t}"
    journal = f"{cont}{vol_issue}{pg}".strip()
    if journal and not journal.endswith("."):
        journal += "."
    return (core + " " + journal + ((" " + tail.strip()) if tail else "")).strip()

# =========================
# Whitelist e validação de citações
# =========================

def extract_candidate_citations_from_text(pages: List[str]) -> set[tuple[str, str]]:
    found: set[tuple[str, str]] = set()
    full_text = "\n".join(pages)
    for pat in CITATION_PATTERNS:
        for m in re.finditer(pat, full_text):
            groups = m.groups()
            if len(groups) == 2:
                surname, year = groups
                found.add((surname.strip(), year.strip()))
            elif len(groups) == 3:
                a, b, year = groups
                found.add((a.strip(), year.strip()))
                found.add((b.strip(), year.strip()))
    return found

def extract_authors_from_frontmatter(pages: List[str]) -> set[str]:
    candidates: set[str] = set()
    head = "\n".join(pages[: min(2, len(pages))])
    lines = _split_nonempty_lines(head)
    for ln in lines:
        if 5 <= len(ln) <= 200 and ("," in ln or " and " in ln or " & " in ln):
            parts = re.split(r",| and | & ", ln)
            for p in parts:
                p = p.strip()
                toks = p.split()
                if len(toks) >= 1:
                    last = toks[-1]
                    if last[:1].isupper() and last.isalpha() and len(last) >= 2:
                        candidates.add(last)
    return candidates

def build_whitelist_for_pdf(pages: List[str]) -> set[tuple[str, str]]:
    wl: set[tuple[str, str]] = set()
    wl |= extract_candidate_citations_from_text(pages)
    for last in extract_authors_from_frontmatter(pages):
        wl.add((last, "*"))
    # Mais a partir das referências
    ref_lines = extract_references_section_lines(pages)
    for block in _split_reference_blocks(ref_lines):
        e = parse_reference_block(block)
        if e.authors and e.year:
            wl.add((e.authors[0].split(",")[0], e.year))
    return wl

def merge_whitelists(list_of_wl: List[set[tuple[str, str]]]) -> set[tuple[str, str]]:
    merged: set[tuple[str, str]] = set()
    for s in list_of_wl:
        merged |= s
    return merged

def validate_generated_citations(text: str, whitelist: set[tuple[str, str]]) -> dict:
    wl_surnames_years = set((s.lower(), y.lower()) for s, y in whitelist)
    wl_surnames_anyyear = set(s.lower() for s, y in whitelist if y == "*")

    found = []
    invalid = []
    valid = []

    patterns = [
        r"\(([A-Z][A-Za-z\-']+),\s*(\d{4}[a-z]?)\)",
        r"\(([A-Z][A-Za-z\-']+)\s+et al\.,\s*(\d{4}[a-z]?)\)",
        r"\(([A-Z][A-Za-z\-']+)\s*&\s*([A-Z][A-Za-z\-']+),\s*(\d{4}[a-z]?)\)",
    ]
    for pat in patterns:
        for m in re.finditer(pat, text):
            groups = m.groups()
            if len(groups) == 2:
                s, y = groups
                s_l, y_l = s.lower(), y.lower()
                if (s_l, y_l) in wl_surnames_years or s_l in wl_surnames_anyyear:
                    valid.append(m.group(0))
                else:
                    invalid.append(m.group(0))
                found.append(m.group(0))
            elif len(groups) == 3:
                a, b, y = groups
                ok_a = (a.lower(), y.lower()) in wl_surnames_years or a.lower() in wl_surnames_anyyear
                ok_b = (b.lower(), y.lower()) in wl_surnames_years or b.lower() in wl_surnames_anyyear
                if ok_a and ok_b:
                    valid.append(m.group(0))
                else:
                    invalid.append(m.group(0))
                found.append(m.group(0))

    summary = f"Citações encontradas: {len(found)} | válidas: {len(valid)} | inválidas: {len(invalid)}"
    return {"invalid": invalid, "valid": valid, "summary": summary}

# =========================
# Construção de prompt
# =========================

def make_prompt(theme: str, selected: List[Chunk], allowed: set[tuple[str, str]] | None = None) -> str:
    """
    Prompt acadêmico com restrição de citações:
    - Cite somente autores/anos presentes em 'allowed'.
    - Use [doc:pag] junto com (Autor, Ano).
    """
    context_blocks = [f"{build_citation(ch)}\n{ch.text}" for ch in selected]
    context = "\n\n".join(context_blocks)

    allowed_str = ""
    if allowed:
        parts = []
        for s, y in sorted(allowed, key=lambda t: (t[0].lower(), t[1])):
            if y == "*":
                parts.append(f"{s} (ano variável)")
            else:
                parts.append(f"{s} ({y})")
        allowed_str = "AUTORIZADOS PARA CITAÇÃO (somente estes): " + "; ".join(parts) + "\n\n"

    return (
        "Você é um pesquisador sênior em marketing/consumo.\n"
        "Responda APENAS com base nos trechos fornecidos; se faltar evidência, diga 'não há suporte nos trechos'.\n"
        "É ESTRITAMENTE PROIBIDO citar autores/anos que NÃO estejam nos PDFs enviados.\n"
        "Use duas formas de citação combinadas: (Autor, Ano) + [doc:pag].\n"
        "Se não for possível citar corretamente, escreva [sem suporte] ao lado da afirmação.\n\n"
        f"{allowed_str}"
        f"TEMA DE PESQUISA: {theme}\n\n"
        "=== TRECHOS (com [doc:pag]) ===\n"
        f"{context}\n\n"
        "=== PRODUZIR (600–1000 palavras) ===\n"
        "1. Introdução ao tema e relevância teórica\n"
        "2. Principais abordagens e contribuições dos autores (com (Autor, Ano) + [doc:pag])\n"
        "3. Síntese e comparação entre perspectivas\n"
        "4. Lacunas e agenda de pesquisa futura (incluir 2–4 proposições testáveis)\n"
        "5. Referências em estilo APA (apenas dos autores/anos autorizados)\n"
    )

# =========================
# LLM providers
# =========================

def call_openai(prompt: str, model_name: str = "gpt-4o-mini") -> str:
    client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))
    try:
        resp = client.chat.completions.create(
            model=model_name,
            messages=[
                {"role": "system", "content": "Você é um assistente de revisão bibliográfica acadêmica e deve se ater ao RAG fornecido."},
                {"role": "user", "content": prompt},
            ],
            temperature=0.4,
        )
        return resp.choices[0].message.content
    except Exception as e:
        st.error(f"Erro ao chamar o modelo OpenAI: {e}")
        return "Falha ao gerar revisão com OpenAI."

def _configure_gemini(api_key: str | None) -> bool:
    if not api_key:
        return False
    try:
        genai.configure(api_key=api_key)
        return True
    except Exception as e:
        st.error(f"Falha ao configurar Gemini: {e}")
        return False

def _list_gemini_models() -> list[str]:
    try:
        models = genai.list_models()
        names = []
        for m in models:
            if hasattr(m, "supported_generation_methods") and "generateContent" in m.supported_generation_methods:
                names.append(getattr(m, "name", None))
        return [n for n in names if n]
    except Exception:
        return []

def _resolve_model_name(requested: str) -> str:
    available = _list_gemini_models()
    if not available:
        return requested
    if requested in available: return requested
    prefixed = f"models/{requested}"
    if prefixed in available: return prefixed
    if requested.startswith("models/") and requested[7:] in available: return requested[7:]
    for cand in ["gemini-1.5-flash", "models/gemini-1.5-flash", "gemini-1.5-pro", "models/gemini-1.5-pro"]:
        if cand in available: return cand
    return available[0]

def call_gemini(prompt: str, model_name: str = "gemini-1.5-flash") -> str:
    try:
        resolved = _resolve_model_name(model_name)
        model = genai.GenerativeModel(resolved)
        resp = model.generate_content(prompt)
        if hasattr(resp, "text") and resp.text:
            return resp.text
        return "Não foi possível gerar texto com o Gemini nesta tentativa (resposta vazia)."
    except Exception as e:
        st.error(f"Falha ao chamar o Gemini: {e}")
        return "Erro ao gerar com Gemini."

# =========================
# Pós-processamento: reconstrução APA a partir dos PDFs
# =========================

def build_pdf_reference_inventory(all_pdf_pages: Dict[str, List[str]]) -> Dict[Tuple[str, str], RefEntry]:
    """
    Constrói um inventário { (first_author_lower, year_lower) -> RefEntry } com base nas seções de referências dos PDFs.
    Deduplica por DOI; se não houver DOI, tenta deduplicar por (autor, ano, título normalizado).
    """
    by_key: Dict[Tuple[str, str], RefEntry] = {}
    seen_doi: Set[str] = set()
    seen_title_sig: Set[Tuple[str, str]] = set()

    for fname, pages in all_pdf_pages.items():
        ref_lines = extract_references_section_lines(pages)
        blocks = _split_reference_blocks(ref_lines)
        for blk in blocks:
            e = parse_reference_block(blk)
            # Deduplicação por DOI
            if e.doi:
                doi_norm = e.doi.lower()
                if doi_norm in seen_doi:
                    continue
                seen_doi.add(doi_norm)
            else:
                # Dedup por (autor, ano, título)
                fa = e.first_author_key() or ""
                y = (e.year or "").lower()
                title_sig = (fa, canonicalize(e.title))
                if (y, title_sig) in seen_title_sig:
                    continue
                seen_title_sig.add((y, title_sig))

            ck = e.citation_key()
            if ck:
                # coloca/atualiza (preferindo entradas com mais metadados)
                prior = by_key.get(ck)
                if (prior is None) or (score_entry(e) > score_entry(prior)):
                    by_key[ck] = e

    return by_key

def score_entry(e: RefEntry) -> int:
    """Heurística simples para preferir entradas mais completas."""
    score = 0
    for v in [e.container, e.volume, e.issue, e.pages, e.doi, e.url, e.title]:
        if v: score += 1
    score += len(e.authors)
    return score

def extract_citation_keys_from_text(text: str) -> Set[Tuple[str, str]]:
    """
    Extrai pares (first_author_lower, year_lower) do texto gerado.
    Trata padrões comuns: (Autor, 2021), (Autor & Autor, 2021), (Autor et al., 2021).
    """
    keys: Set[Tuple[str, str]] = set()
    patterns = [
        r"\(([A-Z][A-Za-z\-']+),\s*(\d{4}[a-z]?)\)",
        r"\(([A-Z][A-Za-z\-']+)\s+et al\.,\s*(\d{4}[a-z]?)\)",
        r"\(([A-Z][A-Za-z\-']+)\s*&\s*([A-Z][A-Za-z\-']+),\s*(\d{4}[a-z]?)\)",
    ]
    for pat in patterns:
        for m in re.finditer(pat, text):
            gs = m.groups()
            if len(gs) == 2:
                s, y = gs
                keys.add((s.lower(), y.lower()))
            elif len(gs) == 3:
                a, b, y = gs
                keys.add((a.lower(), y.lower()))
                keys.add((b.lower(), y.lower()))
    return keys

def build_apa_references_from_output(body_text: str, inventory: Dict[Tuple[str, str], RefEntry]) -> List[str]:
    """
    Monta a seção de referências em APA apenas com base no que foi CITADO no texto
    e no inventário derivado dos PDFs.
    """
    keys_used = extract_citation_keys_from_text(body_text)
    apa_list: List[str] = []
    missing: List[Tuple[str, str]] = []

    for key in sorted(keys_used):
        entry = inventory.get(key)
        if entry:
            apa_list.append(format_apa(entry))
        else:
            missing.append(key)

    if missing:
        # Sinaliza no final quais citações não tiveram match nos PDFs (para auditoria)
        apa_list.append("\n[ATENÇÃO] As seguintes citações não foram encontradas nas referências dos PDFs:")
        apa_list.extend([f"- {k[0].title()}, {k[1]}" for k in missing])

    return apa_list

# =========================
# Fallback sem LLM
# =========================

def extractive_review(theme: str, selected: List[Chunk]) -> str:
    """Fallback sem LLM: apenas concatena trechos + citações."""
    lines = [f"## Revisão (extrativa) — {theme}"]
    for ch in selected:
        lines.append(f"- {ch.text.strip()} {build_citation(ch)}")
    return "\n".join(lines)

# =========================
# Exportação
# =========================

def to_markdown(theme: str, body: str, apa_refs: List[str]) -> str:
    refs_md = "\n".join(f"- {r}" for r in apa_refs) if apa_refs else "*[Sem referências]*"
    return (
        f"# Revisão de Literatura — {theme}\n\n"
        + body
        + "\n\n## Referências (APA)\n"
        + refs_md
    )

def to_docx(theme: str, body_md: str, apa_refs: List[str]) -> bytes:
    doc = Document()
    doc.add_heading(f"Revisão de Literatura — {theme}", level=0)
    # Corpo (simplificado)
    for line in body_md.splitlines():
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("- "):
            doc.add_paragraph(line[2:].strip())
        else:
            if line.strip() == "":
                doc.add_paragraph("")
            else:
                doc.add_paragraph(line)
    # Referências
    doc.add_heading("Referências (APA)", level=1)
    for r in apa_refs:
        doc.add_paragraph(r)
    bio = io.BytesIO()
    doc.save(bio)
    return bio.getvalue()

# =========================
# UI — Streamlit
# =========================

st.set_page_config(page_title="RAG Review — PDFs ➜ Revisão", layout="wide")
st.title("RAG para Revisão Bibliográfica — múltiplos PDFs ➜ Revisão estruturada (anti-alucinação)")

with st.sidebar:
    st.header("Parâmetros")
    chunk_size = st.number_input(
        "Tamanho do chunk (aprox. palavras)",
        min_value=400, max_value=3000, value=int(os.getenv("CHUNK_SIZE", "1200")), step=100
    )
    overlap = st.number_input(
        "Sobreposição", min_value=0, max_value=800, value=int(os.getenv("CHUNK_OVERLAP", "200")), step=50
    )
    topk = st.slider("Top-K inicial (busca)", 5, 100, 30, step=5)
    per_doc = st.slider("Diversidade por documento (máx. chunks)", 1, 10, 3, step=1)
    max_total = st.slider("Máximo total de chunks no prompt", 5, 30, 12, step=1)

    st.subheader("Provedor LLM")
    provider = st.selectbox("Provedor", ["auto (OpenAI→Gemini→Extrativo)", "OpenAI", "Gemini", "Somente extrativo"], index=0)

    st.divider()
    st.subheader("OpenAI")
    openai_key = st.text_input("OPENAI_API_KEY", type="password", value=os.getenv("OPENAI_API_KEY", ""))
    openai_model = st.text_input("OPENAI_MODEL", value=os.getenv("OPENAI_MODEL", "gpt-4o-mini"))  # ou "gpt-5" se tiver acesso
    if openai_key:
        os.environ["OPENAI_API_KEY"] = openai_key
    if openai_model:
        os.environ["OPENAI_MODEL"] = openai_model

    st.subheader("Gemini (fallback)")
    gemini_key = st.text_input("GEMINI_API_KEY", type="password", value=os.getenv("GEMINI_API_KEY", ""))
    gemini_model = st.selectbox("GEMINI_MODEL", ["gemini-1.5-flash", "gemini-1.5-pro"], index=0)
    if USE_GEMINI and gemini_key:
        _configure_gemini(gemini_key)

uploaded = st.file_uploader("Envie múltiplos PDFs", type=["pdf"], accept_multiple_files=True)

st.markdown("---")
col1, col2 = st.columns([2, 1])
with col1:
    theme = st.text_area(
        "Tema / pergunta de pesquisa",
        placeholder="ex.: Impacto de IA generativa na percepção de autenticidade de marcas"
    )
with col2:
    run_btn = st.button("Gerar revisão", type="primary")

# Estado
if "vector" not in st.session_state:
    st.session_state.vector = None
if "mapping" not in st.session_state:
    st.session_state.mapping = {}
if "chunks" not in st.session_state:
    st.session_state.chunks: List[Chunk] = []
if "citation_whitelist" not in st.session_state:
    st.session_state.citation_whitelist: set[tuple[str, str]] = set()
if "pdf_pages" not in st.session_state:
    st.session_state.pdf_pages: Dict[str, List[str]] = {}

# Ingestão em memória a partir dos uploads + construção da whitelist e inventário
if uploaded and st.session_state.vector is None:
    with st.spinner("Processando PDFs, criando índice e inventário de referências..."):
        vec = VectorIndex()
        mapping: Dict[str, str] = {}
        all_chunks: List[Chunk] = []
        all_whitelists: List[set[tuple[str, str]]] = []
        pdf_pages_map: Dict[str, List[str]] = {}

        for i, uf in enumerate(uploaded):
            raw = uf.read()
            pages, title = read_pdf(raw)
            pdf_pages_map[title] = pages

            # whitelist para este PDF
            wl_pdf = build_whitelist_for_pdf(pages)
            all_whitelists.append(wl_pdf)

            doc_id = f"doc{i+1}"
            mapping[os.path.splitext(title)[0]] = title
            cs = chunk_pages(
                pages, title=title, doc_id=doc_id,
                chunk_size=int(chunk_size), overlap=int(overlap)
            )
            all_chunks.extend(cs)

        vec.add(all_chunks)
        st.session_state.vector = vec
        st.session_state.mapping = mapping
        st.session_state.chunks = all_chunks
        st.session_state.citation_whitelist = merge_whitelists(all_whitelists)
        st.session_state.pdf_pages = pdf_pages_map

    st.success(f"Indexados {len(st.session_state.chunks)} chunks de {len(uploaded)} PDFs.")

# Execução da revisão
if run_btn:
    if not theme.strip():
        st.warning("Informe um tema/pergunta de pesquisa.")
    elif st.session_state.vector is None:
        st.warning("Envie os PDFs antes.")
    else:
        with st.spinner("Buscando e sintetizando..."):
            hits = st.session_state.vector.search(theme, k=int(topk))
            selected = diversify_by_doc(hits, per_doc=int(per_doc), max_total=int(max_total))

            with st.expander("Trechos selecionados (para auditoria)"):
                for ch in selected:
                    st.markdown(f"**{build_citation(ch)} — {ch.title}**")
                    st.write(ch.text)
                    st.divider()

            # Prompt com whitelist
            allowed = st.session_state.get("citation_whitelist", set())
            prompt = make_prompt(theme, selected, allowed=allowed)

            # Escolha do provedor
            body_text = ""
            choice = provider.lower()
            use_openai = (("openai" in choice and USE_OPENAI and os.getenv("OPENAI_API_KEY"))
                          or (choice.startswith("auto") and USE_OPENAI and os.getenv("OPENAI_API_KEY")))
            use_gemini = (("gemini" in choice and USE_GEMINI and gemini_key)
                          or (choice.startswith("auto") and (not use_openai) and USE_GEMINI and gemini_key))

            if use_openai:
                model = os.getenv("OPENAI_MODEL", openai_model or "gpt-4o-mini")
                body_text = call_openai(prompt, model_name=model)
            elif use_gemini:
                body_text = call_gemini(prompt, model_name=gemini_model)
            else:
                body_text = extractive_review(theme, selected)

            # Validação de citações
            report = validate_generated_citations(body_text, allowed)

            # Inventário de referências derivado dos PDFs
            inventory = build_pdf_reference_inventory(st.session_state.pdf_pages)

            # Reconstrução APA somente do que foi citado
            apa_refs = build_apa_references_from_output(body_text, inventory)

            st.markdown("### Revisão gerada")
            st.write(body_text)

            with st.expander("Validação de citações (anti-alucinação)"):
                st.write(report["summary"])
                if report["invalid"]:
                    st.error("Citações inválidas detectadas (não presentes nos PDFs):")
                    for item in sorted(set(report["invalid"])):
                        st.write(f"• {item}")
                else:
                    st.success("Nenhuma citação inválida detectada.")

            with st.expander("Referências derivadas dos PDFs (APA)"):
                for r in apa_refs:
                    st.markdown(f"- {r}")

            # Exportação
            st.markdown("---")
            st.subheader("Exportar")
            md = to_markdown(theme, body_text, apa_refs)
            st.download_button(
                "Baixar Markdown (.md)",
                data=md.encode("utf-8"),
                file_name="revisao.md",
                mime="text/markdown"
            )
            if HAS_DOCX:
                docx_bytes = to_docx(theme, body_text, apa_refs)
                st.download_button(
                    "Baixar Word (.docx)",
                    data=docx_bytes,
                    file_name="revisao.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            else:
                st.caption("python-docx não disponível — instale para exportar .docx")

st.markdown(
    """
---
**Dicas**
- Se o PDF for escaneado sem texto, faça OCR antes (ex.: Tesseract).
- Ajuste *Top-K* e *Diversidade por documento* para controlar foco e abrangência.
- Para maior qualidade, troque o embedding para `all-MiniLM-L12-v2` (pode ficar mais pesado).
- Este app não inventa referências: ele reconstrói a lista APA somente a partir do que está nos PDFs e do que foi citado.
"""
)


